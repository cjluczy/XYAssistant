from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, session, send_from_directory, send_file, make_response, render_template_string
import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
import random
import os
import uuid
import io
import pytesseract
from PIL import Image
import pdfplumber
import logging
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 初始化 Flask 应用
app = Flask(__name__)
app.config['SECRET_KEY'] = 'bank-assistant-secret-key-2023'

# 文件上传配置
UPLOAD_FOLDER = os.path.join('static', 'uploads')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 确保上传目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 统一数据库路径至 test/instance/bank_assistant.db（使用绝对路径，确保目录存在）
basedir = os.path.abspath(os.path.dirname(__file__))
instance_dir = os.path.join(basedir, 'instance')
os.makedirs(instance_dir, exist_ok=True)
db_path = os.path.join(instance_dir, 'bank_assistant.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{db_path}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# 初始化数据库
db = SQLAlchemy(app)

# Flask-Login 配置
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'


# 数据库模型
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), default='user')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class BusinessScene(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)
    description = db.Column(db.Text, nullable=False)
    category = db.Column(db.String(50), nullable=False)
    # 元数据：创建与修改人信息
    creator_department = db.Column(db.String(50))
    creator_name = db.Column(db.String(50))
    updater_department = db.Column(db.String(50))
    updater_name = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    steps = db.relationship('SceneStep', backref='scene', lazy=True, cascade='all, delete-orphan')


class SceneStep(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    scene_id = db.Column(db.Integer, db.ForeignKey('business_scene.id'), nullable=False)
    step_number = db.Column(db.Integer, nullable=False)
    description = db.Column(db.Text, nullable=False)
    transaction_code = db.Column(db.String(20))
    details = db.Column(db.Text)
    condition = db.Column(db.String(100))
    image_url = db.Column(db.String(255))

    __table_args__ = (db.UniqueConstraint('scene_id', 'step_number', name='unique_step_number_per_scene'),)


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# 路由定义
@app.route('/')
def index():
    category = request.args.get('category', type=str)
    q = request.args.get('q', type=str)

    query = BusinessScene.query
    if category:
        query = query.filter_by(category=category)
    if q:
        like = f"%{q}%"
        query = query.filter(
            db.or_(
                BusinessScene.name.ilike(like),
                BusinessScene.description.ilike(like),
                BusinessScene.category.ilike(like),
            )
        )

    scenes = query.all()
    return render_template('index.html', scenes=scenes, selected_category=category, q=q)


@app.route('/scene/<int:scene_id>')
def view_scene(scene_id):
    scene = BusinessScene.query.get_or_404(scene_id)
    scene.steps.sort(key=lambda x: x.step_number)
    return render_template('scene_detail.html', scene=scene)


@app.route('/get_step_details/<int:step_id>')
def get_step_details(step_id):
    step = SceneStep.query.get_or_404(step_id)
    return jsonify({
        'description': step.description,
        'transaction_code': step.transaction_code,
        'details': step.details,
        'condition': step.condition
    })


@app.route('/scene/<int:scene_id>/export/pdf')
def export_scene_pdf(scene_id):
    scene = BusinessScene.query.get_or_404(scene_id)
    scene.steps.sort(key=lambda x: x.step_number)
    
    # 创建PDF文件
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                           rightMargin=2*cm, leftMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)
    
    # 创建样式
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CustomHeading1', fontSize=18, 
                            textColor=colors.HexColor('#0b3d91'), 
                            spaceAfter=18, bold=True, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='CustomHeading2', fontSize=16, 
                            textColor=colors.HexColor('#1456b8'), 
                            spaceAfter=12, bold=True))
    styles.add(ParagraphStyle(name='CustomHeading3', fontSize=14, 
                            textColor=colors.HexColor('#2f80ed'), 
                            spaceAfter=8, bold=True))
    styles.add(ParagraphStyle(name='NormalText', fontSize=12, 
                            spaceAfter=6, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name='NoteText', fontSize=12, 
                            textColor=colors.HexColor('#e74c3c'), 
                            spaceAfter=6, alignment=TA_LEFT))
    
    # 创建内容列表
    elements = []
    
    # 添加标题
    elements.append(Paragraph(f"{scene.name}", styles['CustomHeading1']))
    elements.append(Spacer(1, 12))
    
    # 添加描述
    elements.append(Paragraph(f"描述：{scene.description}", styles['NormalText']))
    elements.append(Spacer(1, 6))
    
    # 添加分类
    elements.append(Paragraph(f"业务分类：{scene.category}", styles['NormalText']))
    elements.append(Spacer(1, 6))
    
    # 添加创建信息
    creator_info = f"创建：{scene.creator_department or '-'}{scene.creator_name or '-'} · "
    if scene.created_at:
        creator_info += scene.created_at.strftime('%Y-%m-%d %H:%M')
    elements.append(Paragraph(creator_info, styles['NormalText']))
    
    # 添加更新信息
    if scene.updater_department or scene.updater_name:
        updater_info = f"修改：{scene.updater_department or '-'}{scene.updater_name or '-'} · "
        if scene.updated_at:
            updater_info += scene.updated_at.strftime('%Y-%m-%d %H:%M')
        elements.append(Paragraph(updater_info, styles['NormalText']))
    
    elements.append(Spacer(1, 24))
    
    # 添加操作步骤
    elements.append(Paragraph("操作步骤", styles['CustomHeading2']))
    elements.append(Spacer(1, 12))
    
    # 添加每个步骤
    for step in scene.steps:
        elements.append(Paragraph(f"{step.step_number}. {step.description}", styles['CustomHeading3']))
        
        if step.transaction_code:
            elements.append(Paragraph(f"交易代码：{step.transaction_code}", styles['NormalText']))
        
        if step.details:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("操作说明：", styles['NormalText']))
            elements.append(Paragraph(step.details, styles['NormalText']))
        
        if step.condition:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("执行条件：", styles['NormalText']))
            elements.append(Paragraph(step.condition, styles['NormalText']))
        
        # 如果有图片，尝试添加
        if step.image_url and os.path.exists(os.path.join(app.root_path, 'static', 'uploads', os.path.basename(step.image_url))):
            try:
                img_path = os.path.join(app.root_path, 'static', 'uploads', os.path.basename(step.image_url))
                img = RLImage(img_path, width=5*inch, height=3*inch)
                elements.append(Spacer(1, 12))
                elements.append(img)
            except Exception as e:
                app.logger.error(f"添加图片失败: {e}")
        
        elements.append(Spacer(1, 24))
    
    # 添加风险提示
    elements.append(Paragraph("风险提示", styles['CustomHeading2']))
    elements.append(Spacer(1, 12))
    
    risk_notes = [
        "请严格按照操作步骤执行，确保业务规范",
        "注意客户身份核实，防范冒用风险",
        "仔细检查交易代码和系统返回信息",
        "现金业务必须经过清分和复核",
        "如发现异常情况，立即暂停业务并报告主管"
    ]
    
    for note in risk_notes:
        elements.append(Paragraph(f"• {note}", styles['NoteText']))
    
    # 生成PDF
    doc.build(elements)
    
    # 准备下载
    buffer.seek(0)
    
    # 创建响应
    response = make_response(send_file(buffer, mimetype='application/pdf'))
    response.headers.set('Content-Disposition', 'inline', filename=f"{secure_filename(scene.name)}.pdf")
    
    return response


@app.route('/scene/<int:scene_id>/export/html')
def export_scene_html(scene_id):
    scene = BusinessScene.query.get_or_404(scene_id)
    scene.steps.sort(key=lambda x: x.step_number)
    
    # HTML模板内容
    html_template = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ scene.name }} - 兴业银行柜员SOP助手</title>
    <style>
        body {
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9fcff;
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #0b3d91;
        }
        h1 {
            color: #0b3d91;
            font-size: 24px;
            margin-bottom: 10px;
        }
        h2 {
            color: #1456b8;
            font-size: 20px;
            margin-top: 30px;
            margin-bottom: 15px;
            padding-bottom: 8px;
            border-bottom: 1px solid #ddd;
        }
        h3 {
            color: #2f80ed;
            font-size: 18px;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        .meta-info {
            margin: 10px 0;
            font-size: 14px;
            color: #666;
        }
        .step {
            margin-bottom: 30px;
            padding: 15px;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        }
        .step-details {
            margin-top: 10px;
            padding-left: 20px;
        }
        .step-image {
            margin-top: 15px;
            text-align: center;
        }
        .step-image img {
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
        }
        .risk-notes {
            margin-top: 30px;
            padding: 20px;
            background-color: #fff3f3;
            border-left: 4px solid #e74c3c;
            border-radius: 4px;
        }
        .risk-notes h2 {
            color: #e74c3c;
            border-bottom: none;
            margin-top: 0;
        }
        .risk-note-item {
            margin: 8px 0;
            padding-left: 20px;
            position: relative;
        }
        .risk-note-item:before {
            content: "•";
            position: absolute;
            left: 0;
            color: #e74c3c;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ scene.name }}</h1>
        <div class="meta-info">描述：{{ scene.description }}</div>
        <div class="meta-info">业务分类：{{ scene.category }}</div>
        <div class="meta-info">创建：{{ scene.creator_department or '-' }}{{ scene.creator_name or '-' }} · {% if scene.created_at %}{{ scene.created_at.strftime('%Y-%m-%d %H:%M') }}{% endif %}</div>
        {% if scene.updater_department or scene.updater_name %}
        <div class="meta-info">修改：{{ scene.updater_department or '-' }}{{ scene.updater_name or '-' }} · {% if scene.updated_at %}{{ scene.updated_at.strftime('%Y-%m-%d %H:%M') }}{% endif %}</div>
        {% endif %}
    </div>
    
    <h2>操作步骤</h2>
    
    {% for step in scene.steps %}
    <div class="step">
        <h3>{{ step.step_number }}. {{ step.description }}</h3>
        
        <div class="step-details">
            {% if step.transaction_code %}
            <div><strong>交易代码：</strong>{{ step.transaction_code }}</div>
            {% endif %}
            
            {% if step.details %}
            <div><strong>操作说明：</strong></div>
            <div>{{ step.details }}</div>
            {% endif %}
            
            {% if step.condition %}
            <div><strong>执行条件：</strong>{{ step.condition }}</div>
            {% endif %}
            
            {% if step.image_url %}
            <div class="step-image">
                <img src="{{ step.image_url }}" alt="操作步骤图片" />
            </div>
            {% endif %}
        </div>
    </div>
    {% endfor %}
    
    <div class="risk-notes">
        <h2>风险提示</h2>
        <div class="risk-note-item">请严格按照操作步骤执行，确保业务规范</div>
        <div class="risk-note-item">注意客户身份核实，防范冒用风险</div>
        <div class="risk-note-item">仔细检查交易代码和系统返回信息</div>
        <div class="risk-note-item">现金业务必须经过清分和复核</div>
        <div class="risk-note-item">如发现异常情况，立即暂停业务并报告主管</div>
    </div>
</body>
</html>'''
    
    # 渲染HTML内容
    rendered_html = render_template_string(html_template, scene=scene)
    
    # 创建响应
    response = make_response(rendered_html)
    response.headers.set('Content-Type', 'text/html')
    response.headers.set('Content-Disposition', 'attachment; filename="{}.html"'.format(secure_filename(scene.name)))
    
    return response


@app.route('/scene/<int:scene_id>/export/docx')
def export_scene_docx(scene_id):
    scene = BusinessScene.query.get_or_404(scene_id)
    scene.steps.sort(key=lambda x: x.step_number)
    
    # 创建Word文档
    doc = docx.Document()
    
    # 添加标题
    title = doc.add_heading(scene.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置标题样式
    title_run = title.runs[0]
    title_run.font.name = '微软雅黑'
    title_run.font.color.rgb = RGBColor(11, 61, 145)  # 深蓝色
    
    # 添加元信息
    doc.add_paragraph()  # 空行
    
    meta_info = []
    meta_info.append(f"描述：{scene.description}")
    meta_info.append(f"业务分类：{scene.category}")
    
    if scene.creator_department or scene.creator_name or scene.created_at:
        creator_text = f"创建：{scene.creator_department or '-'}{scene.creator_name or '-'}"
        if scene.created_at:
            creator_text += f" · {scene.created_at.strftime('%Y-%m-%d %H:%M')}"
        meta_info.append(creator_text)
    
    if scene.updater_department or scene.updater_name or scene.updated_at:
        updater_text = f"修改：{scene.updater_department or '-'}{scene.updater_name or '-'}"
        if scene.updated_at:
            updater_text += f" · {scene.updated_at.strftime('%Y-%m-%d %H:%M')}"
        meta_info.append(updater_text)
    
    for meta_line in meta_info:
        p = doc.add_paragraph(meta_line)
        p.runs[0].font.name = '微软雅黑'
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(102, 102, 102)  # 灰色
    
    # 添加操作步骤标题
    doc.add_paragraph()  # 空行
    steps_title = doc.add_heading('操作步骤', level=1)
    steps_title_run = steps_title.runs[0]
    steps_title_run.font.name = '微软雅黑'
    steps_title_run.font.color.rgb = RGBColor(20, 86, 184)  # 中蓝色
    
    # 添加操作步骤
    for step in scene.steps:
        doc.add_paragraph()  # 空行
        
        # 步骤标题
        step_title = doc.add_heading(f"{step.step_number}. {step.description}", level=2)
        step_title_run = step_title.runs[0]
        step_title_run.font.name = '微软雅黑'
        step_title_run.font.color.rgb = RGBColor(47, 128, 237)  # 浅蓝色
        
        # 步骤详情
        details_para = doc.add_paragraph()
        details_run = details_para.add_run()
        details_run.font.name = '微软雅黑'
        details_run.font.size = Pt(11)
        
        details_content = []
        if step.transaction_code:
            details_content.append(f"交易代码：{step.transaction_code}")
        
        if step.details:
            details_content.append(f"操作说明：{step.details}")
        
        if step.condition:
            details_content.append(f"执行条件：{step.condition}")
        
        for detail in details_content:
            details_run.add_text(detail + '\n')
        
        # 添加步骤图片
        if step.image_url:
            try:
                # 尝试从URL添加图片
                import requests
                from PIL import Image
                from io import BytesIO
                
                response = requests.get(step.image_url)
                if response.status_code == 200:
                    img = Image.open(BytesIO(response.content))
                    
                    # 保存图片到临时内存
                    img_buffer = BytesIO()
                    img_format = img.format if img.format else 'PNG'
                    img.save(img_buffer, format=img_format)
                    img_buffer.seek(0)
                    
                    # 添加图片到Word
                    doc.add_picture(img_buffer, width=Inches(5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error adding image: {e}")
                # 添加错误提示文本
                error_para = doc.add_paragraph("图片添加失败")
                error_run = error_para.runs[0]
                error_run.font.color.rgb = RGBColor(231, 76, 60)  # 红色
    
    # 添加风险提示
    doc.add_paragraph()  # 空行
    risk_title = doc.add_heading('风险提示', level=1)
    risk_title_run = risk_title.runs[0]
    risk_title_run.font.name = '微软雅黑'
    risk_title_run.font.color.rgb = RGBColor(231, 76, 60)  # 红色
    
    risk_notes = [
        "请严格按照操作步骤执行，确保业务规范",
        "注意客户身份核实，防范冒用风险",
        "仔细检查交易代码和系统返回信息",
        "现金业务必须经过清分和复核",
        "如发现异常情况，立即暂停业务并报告主管"
    ]
    
    for note in risk_notes:
        p = doc.add_paragraph()
        p.add_run('• ').font.name = '微软雅黑'
        p.add_run(note).font.name = '微软雅黑'
        p.left_indent = Inches(0.5)
    
    # 保存文档到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # 创建响应
    response = make_response(send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
    response.headers.set('Content-Disposition', 'attachment; filename="{}.docx"'.format(secure_filename(scene.name)))
    
    return response


@app.route('/admin/scenes')
@login_required
def admin_scenes():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))

    scenes = BusinessScene.query.all()
    return render_template('admin_scenes.html', scenes=scenes)


@app.route('/admin/scene/new', methods=['GET', 'POST'])
@login_required
def new_scene():
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))

    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        category = request.form.get('category')
        conditions = request.form.getlist('conditions[]')
        image_urls = request.form.getlist('image_urls[]')

        # 检查场景名称是否已存在
        existing_scene = BusinessScene.query.filter_by(name=name).first()
        if existing_scene:
            flash('场景名称已存在，请使用其他名称')
            return render_template('scene_edit.html')

        scene = BusinessScene(
            name=name,
            description=description,
            category=category,
            creator_department=session.get('login_department'),
            creator_name=session.get('login_name'),
            updater_department=session.get('login_department'),
            updater_name=session.get('login_name')
        )
        db.session.add(scene)
        db.session.flush()

        # 处理步骤数据
        steps = request.form.getlist('steps[]')
        transaction_codes = request.form.getlist('transaction_codes[]')
        details = request.form.getlist('details[]')
        # 与条件数组长度对齐
        if conditions and len(conditions) != len(steps):
            # 不抛错，短的用空字符串补齐
            while len(conditions) < len(steps):
                conditions.append('')
        if image_urls and len(image_urls) != len(steps):
            while len(image_urls) < len(steps):
                image_urls.append('')

        for i, (step_desc, trans_code, detail, cond, img) in enumerate(zip(
            steps,
            transaction_codes,
            details,
            conditions or [''] * len(steps),
            image_urls or [''] * len(steps)
        )):
            if step_desc.strip():
                step = SceneStep(
                    scene_id=scene.id,
                    step_number=i + 1,
                    description=step_desc.strip(),
                    transaction_code=trans_code.strip() if trans_code.strip() else None,
                    details=detail.strip() if detail.strip() else None,
                    condition=cond.strip() if cond and cond.strip() else None,
                    image_url=img.strip() if img and img.strip() else None
                )
                db.session.add(step)

        db.session.commit()
        flash('场景创建成功')
        return redirect(url_for('admin_scenes'))

    return render_template('scene_edit.html')


@app.route('/admin/scene/<int:scene_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_scene(scene_id):
    if current_user.role != 'admin':
        flash('权限不足')
        return redirect(url_for('index'))

    scene = BusinessScene.query.get_or_404(scene_id)

    if request.method == 'POST':
        scene.name = request.form.get('name')
        scene.description = request.form.get('description')
        scene.category = request.form.get('category')
        # 记录修改人
        scene.updater_department = session.get('login_department')
        scene.updater_name = session.get('login_name')

        # 删除旧步骤
        SceneStep.query.filter_by(scene_id=scene.id).delete()

        # 添加新步骤
        steps = request.form.getlist('steps[]')
        transaction_codes = request.form.getlist('transaction_codes[]')
        details = request.form.getlist('details[]')
        conditions = request.form.getlist('conditions[]')
        image_urls = request.form.getlist('image_urls[]')
        if conditions and len(conditions) != len(steps):
            while len(conditions) < len(steps):
                conditions.append('')
        if image_urls and len(image_urls) != len(steps):
            while len(image_urls) < len(steps):
                image_urls.append('')

        for i, (step_desc, trans_code, detail, cond, img) in enumerate(zip(
            steps,
            transaction_codes,
            details,
            conditions or [''] * len(steps),
            image_urls or [''] * len(steps)
        )):
            if step_desc.strip():
                step = SceneStep(
                    scene_id=scene.id,
                    step_number=i + 1,
                    description=step_desc.strip(),
                    transaction_code=trans_code.strip() if trans_code.strip() else None,
                    details=detail.strip() if detail.strip() else None,
                    condition=cond.strip() if cond and cond.strip() else None,
                    image_url=img.strip() if img and img.strip() else None
                )
                db.session.add(step)

        db.session.commit()
        flash('场景更新成功')
        return redirect(url_for('admin_scenes'))

    return render_template('scene_edit.html', scene=scene)


@app.route('/admin/scene/<int:scene_id>/delete', methods=['POST'])
@login_required
def delete_scene(scene_id):
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '权限不足'})

    scene = BusinessScene.query.get_or_404(scene_id)
    db.session.delete(scene)
    db.session.commit()

    return jsonify({'success': True, 'message': '场景删除成功'})


@app.route('/login', methods=['GET', 'POST'])
def login():
    FIXED_CAPTCHA = '9527'
    if request.method == 'POST':
        # 新版登录表单字段（部门/NotesID/姓名/密码/验证码）
        department = request.form.get('department')
        notes_id = request.form.get('notes_id')
        name = request.form.get('name')
        password = request.form.get('password')
        captcha = request.form.get('captcha')

        # 简化校验：仅校验固定验证码，其余暂不限制
        if captcha != FIXED_CAPTCHA:
            flash('验证码错误')
            return render_template('login.html', fixed_captcha=FIXED_CAPTCHA)

        # 找到管理员并登录（若不存在则由 init_db 创建）
        admin_user = User.query.filter_by(username='admin').first()
        if not admin_user:
            flash('管理员账户不存在，请重启应用初始化数据库')
            return render_template('login.html', fixed_captcha=FIXED_CAPTCHA)

        # 将登录者信息存入会话，供创建/修改记录使用
        session['login_department'] = department
        session['login_name'] = name
        session['login_notes_id'] = notes_id

        login_user(admin_user)
        flash(f'欢迎 {name or "管理员"}（{department or "未选择部门"}）登录')
        return redirect(url_for('index'))

    return render_template('login.html', fixed_captcha=FIXED_CAPTCHA)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.clear()
    return redirect(url_for('index'))

# 检查文件类型是否允许
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 图片上传路由
@app.route('/upload/image', methods=['POST'])
@login_required
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': '未找到文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    
    if file and allowed_file(file.filename):
        # 生成唯一文件名
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{ext}"
        
        # 保存文件
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        # 返回相对URL供前端使用
        image_url = url_for('static', filename=f'uploads/{unique_filename}')
        return jsonify({'url': image_url, 'filename': unique_filename})
    
    return jsonify({'error': '文件类型不支持'}), 400

# 静态文件访问路由（确保uploads目录可访问）
@app.route('/static/uploads/<path:filename>')
def serve_uploads(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

def perform_pdf_ocr(pdf_path):
    """对PDF文件进行OCR处理"""
    try:
        content = ''
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # 提取页面图像
                img = page.to_image(resolution=300)
                img_bytes = img.original
                # 对图像进行OCR
                text = pytesseract.image_to_string(img_bytes, lang='chi_sim+eng')
                content += text + '\n'
        return content
    except Exception as e:
        return f"OCR处理失败: {str(e)}"

def init_db():
    with app.app_context():
        db.create_all()
        ensure_business_scene_columns()
        ensure_scene_step_columns()

        if not User.query.filter_by(username='admin').first():
            admin_user = User(
                username='admin',
                password_hash=generate_password_hash('admin123'),
                role='admin'
            )
            db.session.add(admin_user)
            db.session.commit()

        if not BusinessScene.query.first():
            add_default_scenes()


def add_default_scenes():
    scenes_data = [
        {
            'name': '本人小额取款',
            'description': '本人小额人民币取款（<5万元）',
            'category': '取款业务',
            'steps': [
                {
                    'description': '收取客户银行卡',
                    'transaction_code': '',
                    'details': '1. 双手接过客户银行卡\n2. 检查银行卡是否完好无损\n3. 确认银行卡为有效状态（非挂失、非过期）\n4. 如为芯片卡，优先使用芯片读取'
                },
                {
                    'description': '进行"三必问"身份核实',
                    'transaction_code': '',
                    'details': '必须完整执行以下三个问题：\n1. 询问客户与账户持有人关系（确认是否为本人）\n2. 询问取款用途（了解资金去向）\n3. 询问是否知晓此笔交易（确认交易意愿）\n注意观察客户回答时的神态和反应'
                },
                {
                    'description': '使用交易码 1159 查卡状态',
                    'transaction_code': '1159',
                    'details': '检查项目包括：\n• 0800 - 证件有效期是否合规\n• 2169 - 凭证状态是否有效（仅显示一条记录）\n• 1160 - 有无管控状态（只收不付、不收不付均不可取款）\n• 2379 - 非柜面限额有无管控\n• 可用余额是否足够（注意是非活期余额）\n• 反洗钱等级是否为次高风险及以上'
                },
                {
                    'description': '使用交易码 2520 个人银行卡取款',
                    'transaction_code': '2520',
                    'details': '1. 输入取款金额，与客户确认\n2. 请客户输入密码\n3. 核对系统返回的交易信息\n4. 执行取款操作'
                },
                {
                    'description': '凭证生成，客户电子签名',
                    'transaction_code': '',
                    'details': '1. 系统生成取款凭证\n2. 指导客户在电子签名设备上签名\n3. 确认签名清晰完整\n4. 如签名不清晰，需重新签署'
                },
                {
                    'description': '提交交易',
                    'transaction_code': '',
                    'details': '1. 确认所有信息无误\n2. 点击提交按钮\n3. 等待系统处理完成\n4. 确认交易成功'
                },
                {
                    'description': '现金清分，交付客户',
                    'transaction_code': '',
                    'details': '1. 从现金箱取出相应金额\n2. 使用点钞机正反两次清点\n3. 手工复核大面额钞票\n4. 加盖业务章（如需要）'
                },
                {
                    'description': '提醒客户确认金额',
                    'transaction_code': '',
                    'details': '1. 双手递出现金并唱票："这是您的X万元，请清点"\n2. 提醒客户当面点清金额\n3. 提示客户注意资金安全\n4. 询问客户是否需要装钞袋'
                }
            ]
        }
    ,
        {
            'name': '本人取款标准',
            'description': '本人取款标准流程（含风控核验与关键交易码）',
            'category': '取款业务',
            'steps': [
                {'description': '收取客户银行卡，确认取款金额。', 'transaction_code': '', 'details': None},
                {'description': '核对客户身份证原件，双人核验身份，并打印身份核查记录。', 'transaction_code': '', 'details': None},
                {'description': '使用交易码 6910 进入【两卡风险核验系统】，查看卡主信息。', 'transaction_code': '6910', 'details': None},
                {'description': '刷卡执行 1159 查卡状态（0800/2169/1160/2379/可用余额/反洗钱等级）。', 'transaction_code': '1159', 'details': None},
                {'description': '执行 2566 查14天内流水，关注当日是否有相近金额入账。', 'transaction_code': '2566', 'details': None},
                {'description': '输入交易码 2520 进入【个人银行卡取款】，先输交易码再点钞。', 'transaction_code': '2520', 'details': None},
                {'description': '客户在柜外清设备上电子签名，签“回单已确认”。', 'transaction_code': '', 'details': None},
                {'description': '扫描身份证复印件等附件，点击【提交】。', 'transaction_code': '', 'details': None},
                {'description': '付出现金需经清分机处理，加盖大黑章/行号章。', 'transaction_code': '', 'details': None},
                {'description': '唱票交付：“{amount}元，请您当面复点确认。”', 'transaction_code': '', 'details': None},
            ]
        },
        {
            'name': '代办取款<5万',
            'description': '代办取款（金额小于5万元）标准流程',
            'category': '取款业务',
            'steps': [
                {'description': '收卡，确认取款金额及代办关系（如“二舅”）。', 'transaction_code': '', 'details': None},
                {'description': '核对代办人身份证原件（卡主身份证可不提供）。', 'transaction_code': '', 'details': None},
                {'description': '使用交易码 6910 查卡主信息。', 'transaction_code': '6910', 'details': None},
                {'description': '执行 1159 查卡状态（不告知卡主余额）。', 'transaction_code': '1159', 'details': None},
                {'description': '执行 2566 查流水（若当日有大额入账需加强审核）。', 'transaction_code': '2566', 'details': None},
                {'description': '执行 2520，选择【是否代理：是】，录入代办人信息（自动联动 2816）。', 'transaction_code': '2520', 'details': '联动 2816 录入代办人信息'},
                {'description': '客户电子签名，扫描代办人身份证，提交。', 'transaction_code': '', 'details': None},
                {'description': '现金清分盖章，唱票交付。', 'transaction_code': '', 'details': None},
            ]
        },
        {
            'name': '代办取款5-20万',
            'description': '代办取款（5万~20万元）标准流程',
            'category': '取款业务',
            'steps': [
                {'description': '收取代办人身份证 + 卡主身份证或户口本（证明关系）。', 'transaction_code': '', 'details': None},
                {'description': '双人核对证件并打印身份核查记录。', 'transaction_code': '', 'details': None},
                {'description': '使用交易码 6910 查卡主信息。', 'transaction_code': '6910', 'details': None},
                {'description': '执行 1159 查卡状态，检查各项。', 'transaction_code': '1159', 'details': None},
                {'description': '执行 2566 查14天流水，关注资金来源（非本人入账需提供依据）。', 'transaction_code': '2566', 'details': None},
                {'description': '执行 2520，选择【是否代理：是】，录入代办人信息（自动联动 2816）。', 'transaction_code': '2520', 'details': '联动 2816 录入代办人信息'},
                {'description': '因金额≥5万，需换人复核。在核查件上写“已换人复核，取{amount}元”，复核人签字。', 'transaction_code': '', 'details': None},
                {'description': '提交后进入远程授权，上传证件及影像。', 'transaction_code': '', 'details': '远程授权'},
                {'description': '客户电子签名，扫描附件，提交。', 'transaction_code': '', 'details': None},
                {'description': '现金清分盖章，唱票交付。', 'transaction_code': '', 'details': None},
            ]
        },
        {
            'name': '代办取款≥20万',
            'description': '代办取款（≥20万元）标准流程（需联系卡主）',
            'category': '取款业务',
            'steps': [
                {'description': '前序步骤同“代办5万~20万”。', 'transaction_code': '', 'details': None},
                {'description': '因金额≥20万，需换人拨打电话联系卡主本人，核实用途/代办人姓名/金额知情。', 'transaction_code': '', 'details': None},
                {'description': '在【代办大额取款台账】登记相关信息。', 'transaction_code': '', 'details': None},
                {'description': '继续执行 2520，提交后远程授权需再次点击【换人复核】。', 'transaction_code': '2520', 'details': '远程授权 + 换人复核'},
                {'description': '后续步骤同上，清分、盖章、唱票交付。', 'transaction_code': '', 'details': None},
            ]
        },
        {
            'name': '外币取款',
            'description': '外币取款标准流程（本人办理）',
            'category': '取款业务',
            'steps': [
                {'description': '外币必须本人办理，核对身份证。', 'transaction_code': '', 'details': None},
                {'description': '使用 6910 查卡主信息。', 'transaction_code': '6910', 'details': None},
                {'description': '执行 1159 查卡状态。', 'transaction_code': '1159', 'details': None},
                {'description': '执行 2561 查外币账户（业务代号 398）。', 'transaction_code': '2561', 'details': '业务代号 398'},
                {'description': '执行 2520 输入金额。', 'transaction_code': '2520', 'details': None},
                {'description': '双人清点外币，打印《提钞业务录入通知书》，客户纸质+电子签名。', 'transaction_code': '', 'details': None},
                {'description': '现金清分盖章，唱票交付。', 'transaction_code': '', 'details': None},
            ]
        }
    ]

    for scene_data in scenes_data:
        # 检查是否已存在该场景
        existing_scene = BusinessScene.query.filter_by(name=scene_data['name']).first()
        if existing_scene:
            # 如果已存在，跳过
            continue

        scene = BusinessScene(
            name=scene_data['name'],
            description=scene_data['description'],
            category=scene_data['category']
        )
        db.session.add(scene)
        db.session.flush()

        for i, step_data in enumerate(scene_data['steps']):
            step = SceneStep(
                scene_id=scene.id,
                step_number=i + 1,
                description=step_data['description'],
                transaction_code=step_data.get('transaction_code'),
                details=step_data.get('details')
            )
            db.session.add(step)

    db.session.commit()
    print("默认场景数据已初始化")


def ensure_business_scene_columns():
    """确保 business_scene 表包含审计字段（SQLite 动态补列）"""
    from sqlalchemy import text
    conn = db.engine.connect()
    try:
        cols = conn.execute(text("PRAGMA table_info('business_scene')")).fetchall()
        existing = {c[1] for c in cols}  # 第二列为列名
        add_statements = []
        if 'creator_department' not in existing:
            add_statements.append("ALTER TABLE business_scene ADD COLUMN creator_department VARCHAR(50)")
        if 'creator_name' not in existing:
            add_statements.append("ALTER TABLE business_scene ADD COLUMN creator_name VARCHAR(50)")
        if 'updater_department' not in existing:
            add_statements.append("ALTER TABLE business_scene ADD COLUMN updater_department VARCHAR(50)")
        if 'updater_name' not in existing:
            add_statements.append("ALTER TABLE business_scene ADD COLUMN updater_name VARCHAR(50)")
        for stmt in add_statements:
            conn.execute(text(stmt))
    finally:
        conn.close()

def ensure_scene_step_columns():
    """确保 scene_step 表包含新字段（SQLite 动态补列）"""
    from sqlalchemy import text
    conn = db.engine.connect()
    try:
        cols = conn.execute(text("PRAGMA table_info('scene_step')")).fetchall()
        existing = {c[1] for c in cols}
        if 'image_url' not in existing:
            conn.execute(text("ALTER TABLE scene_step ADD COLUMN image_url VARCHAR(255)"))
    finally:
        conn.close()


@app.route('/admin/backfill_scene_meta', methods=['POST'])
@login_required
def backfill_scene_meta():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '权限不足'}), 403

    def random_evening_time():
        start_date = datetime(2025, 10, 30)
        end_date = datetime(2025, 11, 11)
        # 随机选择日期
        days = (end_date - start_date).days
        random_day = start_date + timedelta(days=random.randint(0, days))
        # 晚上 22:00 到 23:59:59 之间随机
        hour = random.choice([22, 23])
        minute = random.randint(0, 59)
        second = random.randint(0, 59)
        return random_day.replace(hour=hour, minute=minute, second=second, microsecond=0)

    try:
        scenes = BusinessScene.query.all()
        count = 0
        for s in scenes:
            ts = random_evening_time()
            s.creator_name = '陈中越'
            s.updater_name = '陈中越'
            s.created_at = ts
            s.updated_at = ts
            count += 1
        db.session.commit()
        flash(f'已回填 {count} 个场景：创建/修改人=陈中越，时间为 2025-10-30 至 2025-11-11 间晚间随机值')
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'处理失败: {e}'}), 500

@app.route('/admin/import_defaults', methods=['POST'])
@login_required
def import_defaults():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '权限不足'}), 403
    try:
        add_default_scenes()
        flash('标准场景已导入（按名称去重，已存在的不重复导入）')
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'导入失败: {e}'}), 500


@app.route('/admin/import_from_file', methods=['POST'])
@login_required
def import_from_file():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '权限不足'}), 403
    try:
        name = request.form.get('name', '').strip()
        category = request.form.get('category', '').strip() or '取款业务'
        description = request.form.get('description', '').strip()
        upload = request.files.get('file')
        if not name or not upload:
            return jsonify({'success': False, 'message': '请提供场景名称与文件'}), 400

        # 去重
        if BusinessScene.query.filter_by(name=name).first():
            return jsonify({'success': False, 'message': '同名场景已存在'}), 400

        # 读取文本内容
        filename = upload.filename or ''
        content = ''
        if filename.lower().endswith('.txt'):
            content = upload.read().decode('utf-8', errors='ignore')
        elif filename.lower().endswith('.docx'):
            try:
                from docx import Document  # 需要 python-docx
                temp_path = os.path.join(instance_dir, f'_upload_{datetime.utcnow().timestamp()}.docx')
                upload.save(temp_path)
                doc = Document(temp_path)
                content = '\n'.join(p.text for p in doc.paragraphs)
                os.remove(temp_path)
            except Exception as e:
                return jsonify({'success': False, 'message': f'DOCX解析失败：{e}，可先转换为TXT重试'}), 400
        elif filename.lower().endswith('.pdf'):
            try:
                logger.info(f'开始处理PDF文件: {filename}')
                # 保存临时文件
                temp_path = os.path.join(instance_dir, f'_upload_{datetime.utcnow().timestamp()}.pdf')
                upload.save(temp_path)
                # 使用pdfplumber提取文本
                with pdfplumber.open(temp_path) as pdf:
                    content = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                logger.info(f'PDF文本提取完成，内容长度: {len(content)}')
                # 如果PDF文本提取为空或不完整，尝试OCR
                if not content.strip():
                    logger.info('PDF文本提取结果为空，开始OCR处理')
                    content = perform_pdf_ocr(temp_path)
                # 最后删除临时文件
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception as e:
                logger.error(f'PDF处理失败: {str(e)}')
                return jsonify({'success': False, 'message': f'PDF解析失败：{e}'}), 400
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
            try:
                logger.info(f'开始处理图片文件: {filename}')
                # 读取图片并进行OCR
                image = Image.open(io.BytesIO(upload.read()))
                # 设置中文识别
                content = pytesseract.image_to_string(image, lang='chi_sim+eng')
                logger.info(f'图片OCR完成，内容长度: {len(content)}')
            except Exception as e:
                logger.error(f'图片OCR失败: {str(e)}')
                return jsonify({'success': False, 'message': f'图片OCR失败：{e}'}), 400
        else:
            return jsonify({'success': False, 'message': '仅支持 .txt、.docx、.pdf 或图片文件'}), 400

        # 按编号拆分步骤：匹配“1.”“1、”“1 ”等
        import re
        lines = [line.strip() for line in content.splitlines() if line.strip()]
        steps_text = []
        buffer = ''
        for line in lines:
            if re.match(r'^\s*\d+[\.\、\)]\s+', line):
                if buffer:
                    steps_text.append(buffer.strip())
                buffer = re.sub(r'^\s*\d+[\.\、\)]\s+', '', line)
            else:
                buffer = f'{buffer} {line}'.strip() if buffer else line
        if buffer:
            steps_text.append(buffer.strip())
        if not steps_text:
            # 如果未检测到编号，直接按行作为步骤
            steps_text = lines

        scene = BusinessScene(
            name=name,
            description=description or f'{name}（由文件导入）',
            category=category,
            creator_department=session.get('login_department'),
            creator_name=session.get('login_name'),
            updater_department=session.get('login_department'),
            updater_name=session.get('login_name')
        )
        db.session.add(scene)
        db.session.flush()

        for i, text_line in enumerate(steps_text, 1):
            db.session.add(SceneStep(
                scene_id=scene.id,
                step_number=i,
                description=text_line
            ))
        db.session.commit()
        flash(f'已从文件导入场景：{name}（{len(steps_text)} 个步骤）')
        return jsonify({'success': True, 'count': len(steps_text)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'导入失败: {e}'}), 500

@app.route('/admin/rename_scene_keys', methods=['POST'])
@login_required
def rename_scene_keys():
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '权限不足'}), 403
    # 约定映射：老名称 -> 新缩写
    mapping = {
        'personal_small': '本人小额取款',
        'self_withdrawal_standard': '本人取款标准',
        'proxy_under_5w': '代办取款<5万',
        'proxy_5w_to_20w': '代办取款5-20万',
        'proxy_over_20w': '代办取款≥20万',
        'foreign_currency_withdrawal': '外币取款',
        'WDR_SELF_LT5W': '本人小额取款',
        'WDR_SELF_STD': '本人取款标准',
        'WDR_AGENT_LT5W': '代办取款<5万',
        'WDR_AGENT_5W_20W': '代办取款5-20万',
        'WDR_AGENT_GE20W': '代办取款≥20万',
        'WDR_FOREIGN': '外币取款',
    }
    try:
        changed = 0
        for old, new in mapping.items():
            scene = BusinessScene.query.filter_by(name=old).first()
            if scene:
                # 若新名已存在则跳过，避免冲突
                conflict = BusinessScene.query.filter_by(name=new).first()
                if conflict:
                    continue
                scene.name = new
                changed += 1
        if changed > 0:
            db.session.commit()
        flash(f'重命名完成，共更新 {changed} 个场景标识')
        return jsonify({'success': True, 'changed': changed})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'重命名失败: {e}'}), 500
if __name__ == '__main__':
    init_db()
    app.run(debug=True)